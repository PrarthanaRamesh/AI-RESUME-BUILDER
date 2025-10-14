import streamlit as st
import io
import os
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

# --- Core Functions ---

def configure_api():
    """
    Configures the Google Generative AI API using a local .env file.
    """
    load_dotenv()
    api_key = os.getenv("GOOGLE_API_KEY")

    if not api_key:
        st.error("Google API Key not found.")
        st.info("Please create a .env file in your project folder and add the line: GOOGLE_API_KEY='your-key-here'")
        st.stop()
        
    try:
        genai.configure(api_key=api_key)
        # --- THIS IS THE FINAL, CORRECTED MODEL NAME ---
        return genai.GenerativeModel('models/gemini-pro-latest')
    except Exception as e:
        st.error(f"Failed to configure Google API: {e}")
        st.stop()

def generate_resume_content(model, details):
    """
    Generates the resume content using the Gemini model.
    """
    prompt = f"""
    You are an expert resume writer. Based on the details below, create the content for a professional resume.

    **Important Instructions:**
    - **Do NOT** include the user's name, email, or phone number in your output. This will be added separately.
    - Start directly with the "Summary" section.
    - Use strong, action-oriented verbs and professional language.
    - Format achievements as compelling bullet points.

    **User Details:**
    - Summary Basis: {details['summary']}
    - Work Experience: {details['experience']}
    - Education: {details['education']}
    - Skills: {details['skills']}

    **Required Sections (in this exact order):**
    1. **Summary:** A concise professional summary.
    2. **Experience:** Detailed work history with bulleted achievements.
    3. **Education:** Educational qualifications.
    4. **Skills:** A list of key skills.
    """
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"An error occurred during content generation: {e}")
        return None

def create_docx(details, ai_generated_body):
    """
    Creates a DOCX resume from user details and AI-generated content.
    """
    doc = Document()
    
    # --- Header ---
    p_name = doc.add_paragraph()
    p_name.add_run(details['name']).bold = True
    p_name.runs[0].font.size = Pt(24)
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact_text = f"{details['email']} | {details['phone']}"
    p_contact = doc.add_paragraph()
    p_contact.add_run(contact_text)
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('_' * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- AI-Generated Body ---
    section_headings = ['Summary', 'Experience', 'Education', 'Skills']
    for line in ai_generated_body.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        is_heading = any(line.lower().startswith(h.lower()) for h in section_headings)
        
        if is_heading:
            # Find which heading it is to ensure proper casing
            for heading in section_headings:
                if line.lower().startswith(heading.lower()):
                    p = doc.add_heading(heading, level=1)
                    p.runs[0].font.size = Pt(14)
                    break
        elif line.startswith('*') or line.startswith('-'):
            # Add bullet points, removing the markdown character
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        else:
            doc.add_paragraph(line)
            
    # --- Save to buffer ---
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def build_ui():
    """
    Builds the Streamlit user interface and returns user inputs.
    """
    st.set_page_config(page_title="AI Resume Builder", page_icon="📝", layout="centered")
    st.title("📝 AI Resume Builder")
    st.markdown("Enter your professional details below to generate a polished resume with AI.")

    with st.form("resume_form"):
        st.subheader("Personal Information")
        name = st.text_input("Full Name", placeholder="e.g., Jane Doe")
        email = st.text_input("Email", placeholder="e.g., jane.doe@example.com")
        phone = st.text_input("Phone Number", placeholder="e.g., +91 98765 43210")

        st.subheader("Professional Details")
        summary = st.text_area("Professional Summary", placeholder="A brief summary of your career, skills, and goals.")
        experience = st.text_area("Work Experience", placeholder="List each role on a new line. Include company, role, duration, and key achievements.")
        education = st.text_area("Education", placeholder="List your degrees, college, and graduation year.")
        skills = st.text_area("Key Skills (Comma Separated)", placeholder="e.g., Python, Project Management, Data Analysis")

        submitted = st.form_submit_button("✨ Generate Resume")

    user_details = {
        "name": name, "email": email, "phone": phone,
        "summary": summary, "experience": experience,
        "education": education, "skills": skills
    }
    return submitted, user_details

# --- Main Application Logic ---

def main():
    """
    Main function to run the Streamlit application.
    """
    submitted, user_details = build_ui()
    
    if submitted:
        # Validate essential inputs
        if not all([user_details['name'], user_details['email'], user_details['summary']]):
            st.warning("Please fill in at least Name, Email, and Professional Summary.")
            return

        with st.spinner("🤖 Generating your resume... This may take a moment."):
            model = configure_api()
            ai_content = generate_resume_content(model, user_details)
            
            if ai_content:
                st.subheader("📄 Generated Resume Content")
                st.text_area("Review and edit the content before downloading:", value=ai_content, height=400)

                docx_file = create_docx(user_details, ai_content)
                
                st.download_button(
                    label="📥 Download Resume as DOCX",
                    data=docx_file,
                    file_name=f"{user_details['name'].replace(' ', '_')}_Resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                st.success("Your resume has been generated successfully!")

if __name__ == "__main__":
    main()